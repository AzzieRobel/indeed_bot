from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# ---------- HEADER / NAME ----------
name = doc.add_heading('Your Name', level=0)
name.alignment = WD_ALIGN_PARAGRAPH.CENTER

contact = doc.add_paragraph('Email: your.email@example.com | Phone: (123) 456-7890 | LinkedIn: linkedin.com/in/yourprofile')
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # blank line

# ---------- PROFESSIONAL SUMMARY ----------
doc.add_heading('Professional Summary', level=1)
doc.add_paragraph(
    "Brief summary of your professional experience, skills, and career goals. "
    "Keep it 2-3 sentences."
)

# ---------- SKILLS ----------
doc.add_heading('Skills', level=1)
skills = doc.add_paragraph()
skills.add_run('• Skill 1\n')
skills.add_run('• Skill 2\n')
skills.add_run('• Skill 3\n')
skills.add_run('• Skill 4\n')

# ---------- EXPERIENCE ----------
doc.add_heading('Professional Experience', level=1)

# Job 1
doc.add_heading('Job Title 1 | Company Name', level=2)
doc.add_paragraph('Location | Start Date – End Date')
doc.add_paragraph(
    "• Responsibility or achievement 1\n"
    "• Responsibility or achievement 2\n"
    "• Responsibility or achievement 3"
)

# Job 2
doc.add_heading('Job Title 2 | Company Name', level=2)
doc.add_paragraph('Location | Start Date – End Date')
doc.add_paragraph(
    "• Responsibility or achievement 1\n"
    "• Responsibility or achievement 2\n"
)

# ---------- EDUCATION ----------
doc.add_heading('Education', level=1)
doc.add_heading('Degree | University Name', level=2)
doc.add_paragraph('Location | Graduation Year')
doc.add_paragraph(
    "Relevant coursework, honors, or achievements (optional)."
)

# ---------- FOOTER ----------
section = doc.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "Page "
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------- SAVE DOCUMENT ----------
doc.save('ATS_friendly_resume.docx')

print("ATS-friendly resume created successfully!")
