import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

import google.generativeai as genai
import yaml
from dotenv import load_dotenv
from docx import Document

from _database import Database
from _resume_generator import ResumeCoverLetterGenerator
from _utils import format_job_details_for_summary


load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("Gemini API key required in .env (GEMINI_API_KEY)")

GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
PROXY_HTTP_ADDRESS = os.getenv("PROXY_HTTP_ADDRESS")
PROXY_USERNAME = os.getenv("PROXY_USERNAME")
PROXY_PASSWORD = os.getenv("PROXY_PASSWORD")

RESUME_TEMPLATE_FILENAME = "resume_modern.docx"
COVER_LETTER_TEMPLATE_FILENAME = "cover_letter_standard.docx"
MODEL_FALLBACKS = [
    "gemini-2.0-flash-lite",
    "gemini-flash-lite-latest",
    "gemini-2.0-flash",
    "gemini-flash-latest",
    "gemini-2.0-pro",
    "gemini-pro-latest",
    "gemini-2.5-flash",
    "gemini-2.5-pro",
    "gemini-1.5-pro",
    "gemini-1.0-pro",
    "gemini-pro",
]


def _load_user_profile() -> Dict[str, Any]:
    with open("config.yaml", "r", encoding="utf-8") as config_file:
        config = yaml.safe_load(config_file)
    profile = config.get("user_profile")
    if not profile:
        raise ValueError("user_profile is required in config.yaml")
    return profile


def _build_proxy_uri() -> Optional[str]:
    if not PROXY_HTTP_ADDRESS:
        return None

    if PROXY_USERNAME and PROXY_PASSWORD and "@" not in PROXY_HTTP_ADDRESS:
        if "://" in PROXY_HTTP_ADDRESS:
            scheme, address = PROXY_HTTP_ADDRESS.split("://", 1)
        else:
            scheme, address = "http", PROXY_HTTP_ADDRESS
        return f"{scheme}://{PROXY_USERNAME}:{PROXY_PASSWORD}@{address}"
    return PROXY_HTTP_ADDRESS


def _initialize_gemini_client() -> tuple[genai.GenerativeModel, str, List[str]]:
    proxy_uri = _build_proxy_uri()
    if proxy_uri:
        os.environ["HTTP_PROXY"] = proxy_uri
        os.environ["HTTPS_PROXY"] = proxy_uri

    genai.configure(api_key=GEMINI_API_KEY, transport="rest")
    available = _fetch_available_models()
    model_candidates = _build_model_priority(GEMINI_MODEL, available)

    last_error: Optional[Exception] = None
    for candidate in model_candidates:
        try:
            print(f"Using Gemini model: {candidate}")
            return genai.GenerativeModel(candidate), candidate, model_candidates
        except Exception as exc:
            last_error = exc
            print(f"Warning: Failed to initialize model '{candidate}': {exc}")

    raise RuntimeError(
        "Unable to initialize any Gemini model."
        + (f" Last error: {last_error}" if last_error else "")
    )


def _fetch_available_models() -> Optional[set[str]]:
    try:
        models = list(genai.list_models())
        return {_normalize_model_name(model.name) for model in models}
    except Exception as exc:
        print(f"Warning: Unable to list Gemini models ({exc}).")
        return None


def _build_model_priority(
    preferred_model: str, available_models: Optional[set[str]]
) -> List[str]:
    normalized_preferred = _normalize_model_name(preferred_model)
    ordered = [normalized_preferred] + MODEL_FALLBACKS
    seen = set()
    result: List[str] = []
    for name in ordered:
        normalized = _normalize_model_name(name)
        if normalized in seen:
            continue
        if available_models and normalized not in available_models:
            continue
        result.append(normalized)
        seen.add(normalized)
    if not result:
        result.append(normalized_preferred)
    return result


def _normalize_model_name(model_name: str) -> str:
    if not model_name:
        return model_name
    if "/" in model_name:
        return model_name.split("/")[-1]
    return model_name


class ResumeTestRunner:
    def __init__(self) -> None:
        self.db = Database()
        self.user_profile = _load_user_profile()
        (
            self.gemini_client,
            self.gemini_model_name,
            self.model_candidates,
        ) = _initialize_gemini_client()
        self._model_index = self.model_candidates.index(self.gemini_model_name)

        self.generator = ResumeCoverLetterGenerator()
        self.template_manager = self.generator.template_manager
        self.template_manager.ensure_template_pack()

        self.resume_template = self.template_manager.template_dir / RESUME_TEMPLATE_FILENAME
        self.cover_letter_template = (
            self.template_manager.template_dir / COVER_LETTER_TEMPLATE_FILENAME
        )
        self.output_dir = self.generator.output_dir

    def run(self, num_resumes: int = 1) -> None:
        jobs = self.db.get_jobs_from_db(limit=num_resumes)
        if not jobs:
            print("No jobs found in the database.")
            return

        for idx, job_tuple in enumerate(jobs, start=1):
            job_context = self._build_job_context(job_tuple)
            print(f"\n{'=' * 70}")
            print(f"Generating resume + cover letter for Job #{idx}")
            print(f"Job ID: {job_context['job_id']} | Link: {job_context['job_link']}")
            print(
                f"Company: {job_context['company_name'] or 'N/A'} | "
                f"Location: {job_context['company_address'] or 'N/A'}"
            )
            print("Requesting structured resume plan from Gemini ...")

            try:
                blueprint = self._request_resume_blueprint(
                    job_context["job_description"]
                )
            except Exception as exc:
                print(f"Gemini request failed: {exc}")
                continue

            if not blueprint:
                print("Skipping job due to empty Gemini response.")
                continue

            try:
                artifacts = self._create_documents(job_context, blueprint)
            except Exception as exc:
                print(f"Document generation failed: {exc}")
                continue

            print("Documents created:")
            print(f"  Resume DOCX: {artifacts.get('resume_docx')}")
            print(f"  Cover Letter DOCX: {artifacts.get('cover_letter_docx')}")
            if artifacts.get("resume_pdf"):
                print(f"  Resume PDF: {artifacts['resume_pdf']}")
            if artifacts.get("cover_letter_pdf"):
                print(f"  Cover Letter PDF: {artifacts['cover_letter_pdf']}")

    def _build_job_context(self, job_tuple: Iterable[Any]) -> Dict[str, Any]:
        job_id, job_link, job_summary, job_details_str, _embedding = job_tuple
        job_details = self._safe_load_json(job_details_str)

        company_name = job_details.get("company") or job_details.get("company_name")
        company_address = job_details.get("location")
        job_title = job_details.get("title") or job_details.get("job_title")

        if job_summary:
            job_description = job_summary
        elif job_details:
            job_description = format_job_details_for_summary(job_details)
        else:
            job_description = f"Job ID {job_id}"

        return {
            "job_id": job_id,
            "job_link": job_link,
            "company_name": company_name,
            "company_address": company_address,
            "job_title": job_title,
            "job_description": job_description,
        }

    def _request_resume_blueprint(self, job_description: str) -> Optional[Dict[str, Any]]:
        prompt = f"""
        You are an elite resume writer. Craft a tailored resume blueprint plus a cover letter outline
        using ONLY the provided candidate profile and job summary.
        Respond strictly as JSON with this schema:
        {{
            "resume": {{
                "name": "Full name",
                "role": "Target job title/headline",
                "email": "email@example.com",
                "phone": "optional phone",
                "location": "City, Country",
                "summary": "2-3 sentences tailored summary",
                "key_achievements": ["...", "..."],
                "technical_skills": ["skill", "..."],
                "soft_skills": ["skill", "..."],
                "experience": [
                    {{
                        "title": "",
                        "company": "",
                        "location": "",
                        "duration": "",
                        "highlights": ["bullet", "..."]
                    }}
                ],
                "education": [
                    {{
                        "degree": "",
                        "institution": "",
                        "location": "",
                        "duration": ""
                    }}
                ],
                "certifications": ["optional list"]
            }},
            "cover_letter": {{
                "opening": "Paragraph",
                "body1": "Paragraph",
                "body2": "Paragraph",
                "closing": "Paragraph"
            }}
        }}
        Candidate profile JSON:
        {json.dumps(self.user_profile, indent=2)}

        Job summary:
        {job_description}
        """

        response = self._call_gemini_with_retries(
            prompt,
            generation_config={
                "temperature": 0.4,
                "max_output_tokens": 1600,
            },
        )
        raw_text = self._extract_response_text(response)
        if not raw_text:
            self._log_gemini_debug(response)
            if self._switch_to_next_model():
                print(
                    f"Retrying Gemini request with model '{self.gemini_model_name}'..."
                )
                return self._request_resume_blueprint(job_description)
            return None
        return self._safe_load_json(raw_text)

    def _call_gemini_with_retries(
        self,
        prompt: str,
        **request_kwargs: Any,
    ) -> Any:
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                return self.gemini_client.generate_content(
                    prompt,
                    **request_kwargs,
                )
            except Exception as exc:
                if "429" in str(exc) and attempt < max_attempts - 1:
                    delay = 2 * (attempt + 1)
                    print(f"Gemini rate limit hit. Retrying in {delay}s ...")
                    time.sleep(delay)
                    continue
                raise

    def _create_documents(
        self, job_context: Dict[str, Any], blueprint: Dict[str, Any]
    ) -> Dict[str, Path]:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        resume_data = blueprint.get("resume", {})
        cover_data = blueprint.get("cover_letter", {})

        resume_doc_path = self._render_resume_doc(
            resume_data, job_context, timestamp
        )
        cover_doc_path = self._render_cover_letter_doc(
            resume_data, cover_data, job_context, timestamp
        )

        results: Dict[str, Path] = {
            "resume_docx": resume_doc_path,
            "cover_letter_docx": cover_doc_path,
        }

        return results

    def _render_resume_doc(
        self, resume_data: Dict[str, Any], job_context: Dict[str, Any], timestamp: str
    ) -> Path:
        template_doc = Document(str(self.resume_template))
        replacements = self._build_resume_replacements(resume_data, job_context)
        self._replace_placeholders(template_doc, replacements)

        output_path = (
            self.output_dir / f"resume_{job_context['job_id']}_{timestamp}.docx"
        )
        template_doc.save(output_path)
        return output_path

    def _render_cover_letter_doc(
        self,
        resume_data: Dict[str, Any],
        cover_data: Dict[str, Any],
        job_context: Dict[str, Any],
        timestamp: str,
    ) -> Path:
        template_doc = Document(str(self.cover_letter_template))
        replacements = self._build_cover_letter_replacements(
            resume_data, cover_data, job_context
        )
        self._replace_placeholders(template_doc, replacements)
        output_path = (
            self.output_dir / f"cover_letter_{job_context['job_id']}_{timestamp}.docx"
        )
        template_doc.save(output_path)
        return output_path

    def _build_resume_replacements(
        self, resume_data: Dict[str, Any], job_context: Dict[str, Any]
    ) -> Dict[str, str]:
        contact_parts = [
            resume_data.get("email"),
            resume_data.get("phone"),
            resume_data.get("location"),
        ]
        contact_block = " | ".join([part for part in contact_parts if part])

        replacements = {
            "${FULL_NAME}": resume_data.get("name", self.user_profile.get("name", "")),
            "${TAGLINE}": resume_data.get("role", job_context.get("job_title", "")) or "",
            "${CONTACT_BLOCK}": contact_block or self.user_profile.get("location", ""),
            "${PROFESSIONAL_PROFILE}": resume_data.get("summary", ""),
            "${TECH_STACK}": self._join_list(resume_data.get("technical_skills")),
            "${SOFT_SKILLS}": self._join_list(resume_data.get("soft_skills")),
            "${EXPERIENCE_SECTION}": self._format_experience(resume_data.get("experience")),
            "${PROJECTS_SECTION}": self._format_key_achievements(
                resume_data.get("key_achievements")
            ),
            "${EDUCATION_SECTION}": self._format_education(resume_data.get("education")),
        }
        return replacements

    def _build_cover_letter_replacements(
        self,
        resume_data: Dict[str, Any],
        cover_data: Dict[str, Any],
        job_context: Dict[str, Any],
    ) -> Dict[str, str]:
        return {
            "${DATE}": datetime.now().strftime("%B %d, %Y"),
            "${HIRING_MANAGER}": cover_data.get("hiring_manager", "Hiring Manager"),
            "${COMPANY_NAME}": job_context.get("company_name") or "Company",
            "${COMPANY_ADDRESS}": job_context.get("company_address") or "",
            "${OPENING_PARAGRAPH}": cover_data.get("opening", ""),
            "${BODY_PARAGRAPH_1}": cover_data.get("body1", ""),
            "${BODY_PARAGRAPH_2}": cover_data.get("body2", ""),
            "${CLOSING_PARAGRAPH}": cover_data.get("closing", ""),
            "${SIGN_OFF}": "Sincerely,",
            "${FULL_NAME}": resume_data.get("name", self.user_profile.get("name", "")),
        }

    @staticmethod
    def _replace_placeholders(doc: Document, replacements: Dict[str, str]) -> None:
        def replace_in_paragraphs(paragraphs: Iterable[Any]) -> None:
            for paragraph in paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

        replace_in_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

    @staticmethod
    def _format_experience(experience_items: Optional[List[Dict[str, Any]]]) -> str:
        if not experience_items:
            return ""
        sections: List[str] = []
        for item in experience_items:
            header_parts = [
                item.get("title"),
                item.get("company"),
                item.get("location"),
                item.get("duration"),
            ]
            header = " | ".join([part for part in header_parts if part])
            bullets = item.get("highlights") or item.get("achievements") or []
            bullet_lines = [
                f"• {bullet.strip()}"
                for bullet in bullets
                if isinstance(bullet, str) and bullet.strip()
            ]
            section_text = header
            if bullet_lines:
                section_text += "\n" + "\n".join(bullet_lines)
            sections.append(section_text.strip())
        return "\n\n".join(sections).strip()

    @staticmethod
    def _format_key_achievements(items: Optional[List[str]]) -> str:
        if not items:
            return ""
        return "\n".join(f"• {item.strip()}" for item in items if item and item.strip())

    @staticmethod
    def _format_education(education_items: Optional[List[Dict[str, Any]]]) -> str:
        if not education_items:
            return ""
        entries: List[str] = []
        for edu in education_items:
            parts = [
                edu.get("degree"),
                edu.get("institution"),
                edu.get("location"),
                edu.get("duration") or edu.get("year"),
            ]
            entry = " | ".join([part for part in parts if part])
            if entry:
                entries.append(entry)
        return "\n".join(entries)

    @staticmethod
    def _join_list(items: Optional[List[str]], separator: str = ", ") -> str:
        if not items:
            return ""
        return separator.join([item.strip() for item in items if item and item.strip()])

    @staticmethod
    def _extract_response_text(response: Any) -> str:
        try:
            text = getattr(response, "text", None)
        except Exception:
            text = None
        if text:
            return text

        candidates = getattr(response, "candidates", None)
        if not candidates:
            return ""

        candidate = candidates[0]
        parts_container = getattr(candidate, "content", None)
        part_list = []

        if parts_container is not None:
            nested_parts = getattr(parts_container, "parts", None)
            if nested_parts:
                part_list = nested_parts
            elif isinstance(parts_container, list):
                part_list = parts_container

        if not part_list:
            part_list = getattr(candidate, "parts", []) or []

        collected_text: List[str] = []
        for part in part_list:
            part_text = getattr(part, "text", None)
            if part_text:
                collected_text.append(part_text)

        return "\n".join(collected_text).strip()

    @staticmethod
    def _log_gemini_debug(response: Any) -> None:
        try:
            candidates = getattr(response, "candidates", None)
            if not candidates:
                print("Gemini returned no candidates.")
                return
            candidate = candidates[0]
            finish_reason = getattr(candidate, "finish_reason", None)
            safety = getattr(candidate, "safety_ratings", None)
            print(
                "Gemini finished without content. "
                f"finish_reason={finish_reason}, safety_ratings={safety}"
            )
        except Exception as exc:
            print(f"Gemini debug logging failed: {exc}")

    @staticmethod
    def _safe_load_json(payload: Any) -> Dict[str, Any]:
        if not payload:
            return {}
        if isinstance(payload, dict):
            return payload
        if isinstance(payload, str):
            payload = payload.strip()
            if not payload:
                return {}
            try:
                json_payload = json.loads(payload)
                if isinstance(json_payload, dict):
                    return json_payload
                return {}
            except json.JSONDecodeError:
                # Attempt to extract JSON substring
                json_match = re.search(r"\{[\s\S]*\}", payload)
                if json_match:
                    try:
                        return json.loads(json_match.group())
                    except json.JSONDecodeError:
                        return {}
        return {}

    def _switch_to_next_model(self) -> bool:
        if self._model_index + 1 >= len(self.model_candidates):
            return False
        for idx in range(self._model_index + 1, len(self.model_candidates)):
            candidate = self.model_candidates[idx]
            try:
                self.gemini_client = genai.GenerativeModel(candidate)
                self._model_index = idx
                self.gemini_model_name = candidate
                print(f"Switched Gemini model to '{candidate}'.")
                return True
            except Exception as exc:
                print(f"Warning: Failed to switch to model '{candidate}': {exc}")
        return False

def test_generate_resumes_from_jobs(num_resumes: int = 1) -> None:
    runner = ResumeTestRunner()
    runner.run(num_resumes=num_resumes)


if __name__ == "__main__":
    test_generate_resumes_from_jobs(num_resumes=2)
