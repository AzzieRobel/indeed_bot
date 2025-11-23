"""
Template generation utilities for ATS-friendly resume & cover letter documents.

Phase 1 deliverable: build lightweight template pack with stable placeholder schema.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

TEMPLATE_DIR = Path("assets/templates")


@dataclass(frozen=True)
class TemplateDefinition:
    name: str
    filename: str
    placeholders: Dict[str, str]


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for heading in ("Heading 1", "Heading 2", "Heading 3"):
        h = doc.styles[heading]
        h.font.name = "Calibri"
        h.font.bold = True
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)


def _add_section(doc: Document, title: str, placeholder: str) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(placeholder)


def _add_bullet_placeholder(doc: Document, placeholder_lines: List[str]) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run("\n".join(placeholder_lines))


def _build_resume_classic(path: Path) -> None:
    doc = Document()
    _set_base_style(doc)

    header = doc.add_heading("${FULL_NAME}", level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("${PROFESSIONAL_TITLE}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = doc.add_paragraph("${CONTACT_BLOCK}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("${SUMMARY}")

    doc.add_heading("Skills", level=1)
    _add_bullet_placeholder(
        doc,
        [
            "• ${CORE_SKILL_1}",
            "• ${CORE_SKILL_2}",
            "• ${CORE_SKILL_3}",
            "• ${CORE_SKILL_4}",
        ],
    )

    _add_section(doc, "Experience", "${EXPERIENCE_SECTION}")
    _add_section(doc, "Education", "${EDUCATION_SECTION}")
    _add_section(doc, "Certifications", "${CERTIFICATIONS_SECTION}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _build_resume_modern(path: Path) -> None:
    doc = Document()
    _set_base_style(doc)

    name = doc.add_heading("${FULL_NAME}", level=0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline = doc.add_paragraph("${TAGLINE}")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = doc.add_paragraph("${CONTACT_BLOCK}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("Professional Profile", level=1)
    doc.add_paragraph("${PROFESSIONAL_PROFILE}")

    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("${TECH_STACK}")

    doc.add_heading("Soft Skills", level=1)
    doc.add_paragraph("${SOFT_SKILLS}")

    _add_section(doc, "Experience", "${EXPERIENCE_SECTION}")
    _add_section(doc, "Projects", "${PROJECTS_SECTION}")
    _add_section(doc, "Education", "${EDUCATION_SECTION}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _build_cover_letter_standard(path: Path) -> None:
    doc = Document()
    _set_base_style(doc)

    doc.add_paragraph("${DATE}")
    doc.add_paragraph("${HIRING_MANAGER}")
    doc.add_paragraph("${COMPANY_NAME}")
    doc.add_paragraph("${COMPANY_ADDRESS}")

    doc.add_paragraph("${OPENING_PARAGRAPH}")
    doc.add_paragraph("${BODY_PARAGRAPH_1}")
    doc.add_paragraph("${BODY_PARAGRAPH_2}")
    doc.add_paragraph("${CLOSING_PARAGRAPH}")
    closing = doc.add_paragraph()
    closing.add_run("${SIGN_OFF}")
    closing.add_run("\n${FULL_NAME}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


TEMPLATES: List[TemplateDefinition] = [
    TemplateDefinition(
        name="Modern Resume",
        filename="resume_modern.docx",
        placeholders={
            "FULL_NAME": "Candidate full name",
            "TAGLINE": "Short value proposition",
            "CONTACT_BLOCK": "Email, phone, location",
            "PROFESSIONAL_PROFILE": "Professional profile paragraph",
            "TECH_STACK": "Technical stack keywords",
            "SOFT_SKILLS": "Soft skills list",
            "EXPERIENCE_SECTION": "Experience entries",
            "PROJECTS_SECTION": "Projects/achievements",
            "EDUCATION_SECTION": "Education information",
        },
    ),
    TemplateDefinition(
        name="Standard Cover Letter",
        filename="cover_letter_standard.docx",
        placeholders={
            "DATE": "Current date",
            "HIRING_MANAGER": "Hiring manager name",
            "COMPANY_NAME": "Company name",
            "COMPANY_ADDRESS": "Company address",
            "OPENING_PARAGRAPH": "Opening paragraph",
            "BODY_PARAGRAPH_1": "Body paragraph emphasizing fit",
            "BODY_PARAGRAPH_2": "Optional second body paragraph",
            "CLOSING_PARAGRAPH": "Closing/CTA paragraph",
            "SIGN_OFF": "Closing phrase",
            "FULL_NAME": "Candidate name",
        },
    ),
]


class TemplateManager:
    def __init__(self, template_dir: Path = TEMPLATE_DIR) -> None:
        self.template_dir = template_dir

    def ensure_template_pack(self, overwrite: bool = False) -> None:
        builders = {
            "resume_classic.docx": _build_resume_classic,
            "resume_modern.docx": _build_resume_modern,
            "cover_letter_standard.docx": _build_cover_letter_standard,
        }
        self.template_dir.mkdir(parents=True, exist_ok=True)
        for definition in TEMPLATES:
            output_path = self.template_dir / definition.filename
            if output_path.exists() and not overwrite:
                continue
            builders[definition.filename](output_path)

    def placeholder_schema(self) -> Dict[str, Dict[str, str]]:
        return {definition.filename: definition.placeholders for definition in TEMPLATES}


if __name__ == "__main__":
    TemplateManager().ensure_template_pack()
    print(f"Templates written to {TEMPLATE_DIR.resolve()}")
